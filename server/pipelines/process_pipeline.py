from __future__ import annotations

import asyncio
import inspect
import logging
import mimetypes
import re
import time
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from uuid import uuid4
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from server.clients.ollama_client import OllamaServiceError
from server.document_registry import DocumentRegistry
from server.pipelines.classifier import ClassificationValidationError, DocumentClassifier
from server.pipelines.extractor import DocumentExtractor, ExtractionValidationError
from server.pipelines.file_organizer import FileOrganizer
from server.pipelines.noop_organizer import NoOpOrganizer
from server.pipelines.search import IndexedDocument
from server.pipelines.thumbnails import generate_thumbnail
from server.pipelines.whisper_proxy import WhisperProxy, WhisperProxyError
from server.schemas import (
    DocumentClassification,
    ExtractionResult,
    MoveExecutor,
    MovePlan,
    MoveResult,
    ProcessResponse,
    ProcessDiagnostics,
    SourceModality,
    TranscriptionResponse,
    UiDocumentKind,
    UiDocumentRecord,
)

SUPPORTED_TEXT_TYPES = {
    "text/plain",
    "text/markdown",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
SUPPORTED_IMAGE_TYPES = {
    "image/jpeg",
    "image/png",
    "image/webp",
}
SUPPORTED_AUDIO_TYPES = {
    "audio/mpeg",
    "audio/wav",
    "audio/x-wav",
    "audio/mp4",
    "audio/m4a",
    "audio/aiff",
    "audio/x-aiff",
}


class UnsupportedMediaTypeError(ValueError):
    """Raised for file types that the current processing pipeline does not support."""


logger = logging.getLogger(__name__)


class DocumentProcessPipeline:
    def __init__(
        self,
        *,
        classifier: DocumentClassifier,
        extractor: DocumentExtractor,
        organizer: FileOrganizer | NoOpOrganizer,
        whisper_service: WhisperProxy | None = None,
        document_registry: DocumentRegistry | None = None,
        realtime_manager: object | None = None,
        search_pipeline: object | None = None,
        max_text_characters: int = 12000,
        classifier_max_text_characters: int | None = None,
        llm_sequence_lock: asyncio.Lock | None = None,
    ) -> None:
        self.classifier = classifier
        self.extractor = extractor
        self.organizer = organizer
        self.whisper_service = whisper_service
        self.document_registry = document_registry
        self.realtime_manager = realtime_manager
        self.search_pipeline = search_pipeline
        self.max_text_characters = max_text_characters
        self.classifier_max_text_characters = classifier_max_text_characters or min(max_text_characters, 4000)
        self._background_tasks: set[asyncio.Task[None]] = set()
        self._llm_sequence_lock = llm_sequence_lock or asyncio.Lock()

    async def process_upload(
        self,
        *,
        filename: str,
        content: bytes,
        content_type: str | None,
        execute_move: bool,
        source_path: str | None,
        client_id: str | None = None,
        client_request_id: str | None = None,
        move_executor: MoveExecutor = "none",
    ) -> ProcessResponse:
        request_id = client_request_id or str(uuid4())
        record_id = str(uuid4())
        created_at = datetime.now(UTC).isoformat()
        mime_type = self._detect_mime(filename, content_type)
        source_modality = self._detect_modality(mime_type)
        thumbnail_data: str | None = generate_thumbnail(content, mime_type)
        timings: dict[str, float] = {}
        errors: list[str] = []
        warnings: list[str] = []
        pipeline_flags: list[str] = []
        transcription: TranscriptionResponse | None = None
        error_code: str | None = None
        retryable = False
        fallback_reason: str | None = None
        classifier_raw_response_path: str | None = None
        self._log_pipeline_event(
            "pipeline.received",
            request_id=request_id,
            client_id=client_id,
            filename=filename,
            mime_type=mime_type,
            source_modality=source_modality,
            record_id=record_id,
        )

        await self._emit_event(
            client_id,
            {
                "type": "job.started",
                "request_id": request_id,
                "client_id": client_id,
                "job_kind": "process",
                "filename": filename,
                "source_modality": source_modality,
                "thumbnail_data": thumbnail_data,
            },
        )
        await self._progress(client_id, request_id, "processing", "Bearbetar dokument")

        try:
            extracted_text: str
            pdf_image_fallback: tuple[bytes, str] | None = None
            if mime_type in SUPPORTED_IMAGE_TYPES:
                extracted_text = ""
            elif mime_type in SUPPORTED_TEXT_TYPES:
                extracted_text = self._extract_text(content, mime_type)
                if mime_type == "application/pdf" and not extracted_text.strip():
                    pdf_image_fallback = self._extract_pdf_image_for_classification(content)
            elif mime_type in SUPPORTED_AUDIO_TYPES:
                extracted_text, transcription, error_code, retryable = await self._process_audio(
                    filename=filename,
                    content=content,
                    mime_type=mime_type,
                    client_id=client_id,
                    request_id=request_id,
                )
                if transcription is None:
                    classification = self._fallback_classification(
                        filename=filename,
                        extracted_text="",
                        source_modality=source_modality,
                        source_path=source_path,
                    )
                    pipeline_flags.append("audio_processing_unavailable")
                    extraction = ExtractionResult(fields={}, field_confidence={}, missing_fields=[])
                    response = ProcessResponse(
                        request_id=request_id,
                        status="failed_runtime",
                        mime_type=mime_type,
                        classification=classification,
                        extraction=extraction,
                        move_plan=MovePlan(reason="no_matching_rule"),
                        move_result=MoveResult(from_path=source_path),
                        timings=timings,
                        errors=[error_code or "audio_processing_unavailable"],
                        record_id=record_id,
                        source_modality=source_modality,
                        created_at=created_at,
                        transcription=None,
                        ui_kind="audio",
                        undo_token=None,
                        move_status="not_requested",
                        retryable=retryable,
                        error_code=error_code or "audio_processing_unavailable",
                        warnings=["Audio processing unavailable."],
                        diagnostics=ProcessDiagnostics(
                            pipeline_flags=pipeline_flags,
                            fallback_reason="audio_processing_unavailable",
                        ),
                        thumbnail_data=thumbnail_data,
                    )
                    self._persist_record(response)
                    await self._emit_event(
                        client_id,
                        {
                            "type": "job.failed",
                            "request_id": request_id,
                            "client_id": client_id,
                            "message": error_code or "audio_processing_unavailable",
                        },
                    )
                    return response
            else:
                raise UnsupportedMediaTypeError(mime_type)

            await self._progress(client_id, request_id, "processing", "Väntar på modellkön")
            llm_wait_started = time.perf_counter()
            self._log_pipeline_event(
                "pipeline.llm.waiting",
                request_id=request_id,
                client_id=client_id,
                filename=filename,
                mime_type=mime_type,
                source_modality=source_modality,
                record_id=record_id,
            )
            async with self._llm_sequence_lock:
                self._log_pipeline_event(
                    "pipeline.llm.acquired",
                    request_id=request_id,
                    client_id=client_id,
                    filename=filename,
                    mime_type=mime_type,
                    source_modality=source_modality,
                    record_id=record_id,
                    wait_ms=round((time.perf_counter() - llm_wait_started) * 1000, 2),
                )
                classify_started = time.perf_counter()
                self._log_pipeline_event(
                    "pipeline.classify.start",
                    request_id=request_id,
                    client_id=client_id,
                    filename=filename,
                    mime_type=mime_type,
                    source_modality=source_modality,
                    record_id=record_id,
                )
                if mime_type in SUPPORTED_IMAGE_TYPES:
                    classification, used_fallback, fallback_reason, classifier_raw_response_path = await self._classify_image(
                        content,
                        mime_type,
                        request_id,
                        filename=filename,
                        source_path=source_path,
                    )
                    extracted_text = classification.ocr_text or classification.summary
                elif pdf_image_fallback is not None:
                    fallback_image_bytes, fallback_mime = pdf_image_fallback
                    classification, used_fallback, fallback_reason, classifier_raw_response_path = await self._classify_image(
                        fallback_image_bytes,
                        fallback_mime,
                        request_id,
                        filename=filename,
                        source_path=source_path,
                    )
                    extracted_text = classification.ocr_text or classification.summary
                    pipeline_flags.append("pdf_text_empty_image_fallback")
                else:
                    classification, used_fallback, fallback_reason, classifier_raw_response_path = await self._classify_text(
                        extracted_text,
                        request_id,
                        filename=filename,
                        source_path=source_path,
                    )

                if self._classification_has_empty_core_fields(classification):
                    classification = self._fallback_classification(
                        filename=filename,
                        extracted_text=extracted_text,
                        source_modality=source_modality,
                        source_path=source_path,
                    )
                    used_fallback = True
                    fallback_reason = "classifier_empty_fields"
                    pipeline_flags.append("classifier_empty_fields_fallback")

                timings["classify_ms"] = round((time.perf_counter() - classify_started) * 1000, 2)
                self._log_pipeline_event(
                    "pipeline.classify.done",
                    request_id=request_id,
                    client_id=client_id,
                    filename=filename,
                    mime_type=mime_type,
                    source_modality=source_modality,
                    record_id=record_id,
                    elapsed_ms=timings["classify_ms"],
                )
                if used_fallback:
                    error_code = "classification_validation_fallback"
                    if fallback_reason == "classifier_invalid_json":
                        pipeline_flags.append("classifier_invalid_json_fallback")
                    warnings.append("Kunde inte tolka dokumentet fullt ut, visning sker i generiskt läge.")
                    self._log_pipeline_event(
                        "pipeline.classify.fallback",
                        request_id=request_id,
                        client_id=client_id,
                        filename=filename,
                        mime_type=mime_type,
                        source_modality=source_modality,
                        record_id=record_id,
                        fallback_reason=fallback_reason,
                        raw_response_path=classifier_raw_response_path,
                    )
                await self._progress(
                    client_id, request_id, "classified", "Dokument klassificerat",
                    data={"classification": classification.model_dump(mode="json")},
                )

                if self._should_skip_extraction(classification.document_type):
                    extraction = ExtractionResult(fields={}, field_confidence={}, missing_fields=[])
                    timings["extract_ms"] = 0.0
                    self._log_pipeline_event(
                        "pipeline.extract.skipped",
                        request_id=request_id,
                        client_id=client_id,
                        filename=filename,
                        mime_type=mime_type,
                        source_modality=source_modality,
                        record_id=record_id,
                        document_type=classification.document_type,
                    )
                else:
                    await self._progress(client_id, request_id, "extracting", "Extraherar fält")
                    extract_started = time.perf_counter()
                    self._log_pipeline_event(
                        "pipeline.extract.start",
                        request_id=request_id,
                        client_id=client_id,
                        filename=filename,
                        mime_type=mime_type,
                        source_modality=source_modality,
                        record_id=record_id,
                    )
                    try:
                        extraction = await self.extractor.extract(
                            extracted_text[: self.max_text_characters],
                            classification,
                            request_id=request_id,
                        )
                    except ExtractionValidationError:
                        extraction = ExtractionResult(fields={}, field_confidence={}, missing_fields=[])
                        pipeline_flags.append("extractor_invalid_json_fallback")
                    timings["extract_ms"] = round((time.perf_counter() - extract_started) * 1000, 2)
                    self._log_pipeline_event(
                        "pipeline.extract.done",
                        request_id=request_id,
                        client_id=client_id,
                        filename=filename,
                        mime_type=mime_type,
                        source_modality=source_modality,
                        record_id=record_id,
                        elapsed_ms=timings["extract_ms"],
                    )

            await self._progress(
                client_id, request_id, "extracted", "Fält extraherade",
                data={"extraction": extraction.model_dump(mode="json")},
            )
            await self._progress(client_id, request_id, "organizing", "Planerar filsortering")
            plan_started = time.perf_counter()
            self._log_pipeline_event(
                "pipeline.organize.start",
                request_id=request_id,
                client_id=client_id,
                filename=filename,
                mime_type=mime_type,
                source_modality=source_modality,
                record_id=record_id,
            )
            move_plan = self.organizer.plan_move(filename, classification)
            move_result = MoveResult(
                attempted=False,
                success=False,
                from_path=source_path,
                to_path=None,
                error=None,
            )
            move_status = "not_requested"
            status = "move_planned"
            undo_token: str | None = None

            if move_plan.destination:
                if move_plan.auto_move_allowed and move_executor == "client":
                    move_status = "auto_pending_client"
                elif move_plan.auto_move_allowed:
                    move_status = "planned"
                else:
                    move_status = "awaiting_confirmation"
                    await self._progress(
                        client_id,
                        request_id,
                        "awaiting_confirmation",
                        "Väntar på bekräftelse för filflytt",
                    )

            timings["organize_ms"] = round((time.perf_counter() - plan_started) * 1000, 2)
            self._log_pipeline_event(
                "pipeline.organize.done",
                request_id=request_id,
                client_id=client_id,
                filename=filename,
                mime_type=mime_type,
                source_modality=source_modality,
                record_id=record_id,
                elapsed_ms=timings["organize_ms"],
                move_status=move_status,
            )

            ui_kind = self._resolve_ui_kind(
                document_type=classification.document_type,
                source_modality=source_modality,
            )
            if move_status == "awaiting_confirmation":
                status = "move_planned"
            response = ProcessResponse(
                request_id=request_id,
                status=status,
                mime_type=mime_type,
                classification=classification,
                extraction=extraction,
                move_plan=self._coerce_move_reason(move_plan, execute_move, move_executor),
                move_result=move_result,
                timings=timings,
                errors=errors,
                record_id=record_id,
                source_modality=source_modality,
                created_at=created_at,
                transcription=transcription,
                ui_kind=ui_kind,
                undo_token=undo_token,
                move_status=move_status,
                retryable=retryable,
                error_code=error_code,
                warnings=warnings,
                diagnostics=ProcessDiagnostics(
                    pipeline_flags=pipeline_flags,
                    classifier_raw_response_path=classifier_raw_response_path,
                    fallback_reason=fallback_reason,
                ),
                thumbnail_data=thumbnail_data,
            )
            self._persist_record(response)
            self._log_pipeline_event(
                "pipeline.persist.done",
                request_id=request_id,
                client_id=client_id,
                filename=filename,
                mime_type=mime_type,
                source_modality=source_modality,
                record_id=record_id,
                move_status=move_status,
            )

            if undo_token is not None and move_result.from_path and move_result.to_path:
                await self._emit_event(
                    client_id,
                    {
                        "type": "file.moved",
                        "request_id": request_id,
                        "client_id": client_id,
                        "record_id": record_id,
                        "from_path": move_result.from_path,
                        "to_path": move_result.to_path,
                        "undo_token": undo_token,
                    },
                )
                await self._progress(client_id, request_id, "moved", "Filen flyttades")

            if self.search_pipeline is not None:
                indexed_source_path = move_result.to_path or source_path or filename
                self._schedule_indexing(
                    request_id=request_id,
                    client_id=client_id,
                    filename=filename,
                    mime_type=mime_type,
                    source_modality=source_modality,
                    record_id=record_id,
                    ui_kind=ui_kind,
                    move_status=move_status,
                    indexed_document=IndexedDocument(
                        doc_id=record_id,
                        title=classification.title,
                        source_path=indexed_source_path,
                        text=extracted_text,
                        metadata={
                            "document_type": classification.document_type,
                            "summary": classification.summary,
                            "tags": classification.tags,
                        },
                    ),
                )
            elif move_status not in {"auto_pending_client", "awaiting_confirmation"}:
                await self._emit_completed_event(
                    client_id=client_id,
                    request_id=request_id,
                    record_id=record_id,
                    ui_kind=ui_kind,
                )
            self._log_pipeline_event(
                "pipeline.response.ready",
                request_id=request_id,
                client_id=client_id,
                filename=filename,
                mime_type=mime_type,
                source_modality=source_modality,
                record_id=record_id,
                move_status=move_status,
            )
            return response
        except Exception as error:
            self._log_pipeline_event(
                "pipeline.failed",
                request_id=request_id,
                client_id=client_id,
                filename=filename,
                mime_type=mime_type,
                source_modality=source_modality,
                record_id=record_id,
                error=str(error),
            )
            await self._emit_event(
                client_id,
                {
                    "type": "job.failed",
                    "request_id": request_id,
                    "client_id": client_id,
                    "message": str(error),
                },
            )
            raise

    async def _classify_image(
        self,
        content: bytes,
        mime_type: str,
        request_id: str,
        *,
        filename: str,
        source_path: str | None,
    ) -> tuple[DocumentClassification, bool, str | None, str | None]:
        try:
            return await self.classifier.classify_image(content, mime_type, request_id=request_id), False, None, None
        except ClassificationValidationError as error:
            return (
                self._fallback_classification(
                    filename=filename,
                    extracted_text="",
                    source_modality="image",
                    source_path=source_path,
                ),
                True,
                "classifier_invalid_json",
                error.raw_response_path,
            )

    async def _classify_text(
        self,
        text: str,
        request_id: str,
        *,
        filename: str,
        source_path: str | None,
    ) -> tuple[DocumentClassification, bool, str | None, str | None]:
        try:
            return await self.classifier.classify_text(
                text[: self.classifier_max_text_characters],
                request_id=request_id,
            ), False, None, None
        except ClassificationValidationError as error:
            return (
                self._fallback_classification(
                    filename=filename,
                    extracted_text=text,
                    source_modality="text",
                    source_path=source_path,
                ),
                True,
                "classifier_invalid_json",
                error.raw_response_path,
            )

    async def _process_audio(
        self,
        *,
        filename: str,
        content: bytes,
        mime_type: str,
        client_id: str | None,
        request_id: str,
    ) -> tuple[str, TranscriptionResponse | None, str | None, bool]:
        if self.whisper_service is None:
            return "", None, "audio_processing_unavailable", True
        await self._progress(client_id, request_id, "transcribing", "Transkriberar ljud")
        try:
            transcription = await self.whisper_service.transcribe(
                filename=filename,
                content=content,
                content_type=mime_type,
                client_id=client_id,
                client_request_id=request_id,
            )
            return transcription.text, transcription, None, False
        except WhisperProxyError as error:
            retryable = error.status_code >= 500
            return "", None, "audio_processing_unavailable", retryable

    def _persist_record(self, response: ProcessResponse) -> None:
        if self.document_registry is None:
            return
        self.document_registry.upsert_document(
            UiDocumentRecord(
                id=response.record_id or str(uuid4()),
                request_id=response.request_id,
                title=response.classification.title,
                summary=response.classification.summary,
                mime_type=response.mime_type,
                source_modality=response.source_modality or "text",
                kind=response.ui_kind or "generic",
                document_type=response.classification.document_type,
                template=response.classification.template,
                source_path=response.move_result.to_path or response.move_result.from_path,
                created_at=response.created_at or datetime.now(UTC).isoformat(),
                updated_at=datetime.now(UTC).isoformat(),
                classification=response.classification,
                extraction=response.extraction,
                transcription=response.transcription,
                move_plan=response.move_plan,
                move_result=response.move_result,
                tags=response.classification.tags,
                status="completed" if response.status in {"move_planned", "move_executed"} and response.move_status not in {"awaiting_confirmation", "auto_pending_client"} else response.status,
                undo_token=response.undo_token,
                move_status=response.move_status,
                retryable=response.retryable,
                error_code=response.error_code,
                warnings=response.warnings,
                diagnostics=response.diagnostics,
                thumbnail_data=response.thumbnail_data,
            )
        )

    async def drain_background_tasks(self) -> None:
        if not self._background_tasks:
            return
        await asyncio.gather(*list(self._background_tasks), return_exceptions=True)

    @staticmethod
    def _detect_mime(filename: str, content_type: str | None) -> str:
        if content_type:
            return content_type
        guessed, _ = mimetypes.guess_type(filename)
        return guessed or "application/octet-stream"

    @staticmethod
    def _extract_text(content: bytes, mime_type: str) -> str:
        if mime_type in {"text/plain", "text/markdown"}:
            return content.decode("utf-8")
        if mime_type == "application/pdf":
            reader = PdfReader(BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            document = DocxDocument(BytesIO(content))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        raise UnsupportedMediaTypeError(mime_type)

    @staticmethod
    def _extract_pdf_image_for_classification(content: bytes) -> tuple[bytes, str] | None:
        try:
            reader = PdfReader(BytesIO(content))
        except Exception:  # pragma: no cover - guarded by parser behavior in _extract_text
            return None
        for page in reader.pages:
            images = getattr(page, "images", [])
            for image in images:
                image_bytes = getattr(image, "data", None)
                if not isinstance(image_bytes, (bytes, bytearray)):
                    continue
                image_mime = DocumentProcessPipeline._infer_image_mime_from_pdf_image(
                    image_name=getattr(image, "name", None),
                    image_bytes=bytes(image_bytes),
                )
                if image_mime not in SUPPORTED_IMAGE_TYPES:
                    continue
                return bytes(image_bytes), image_mime
        return None

    @staticmethod
    def _infer_image_mime_from_pdf_image(*, image_name: Any, image_bytes: bytes) -> str | None:
        if isinstance(image_name, str):
            guessed_mime, _ = mimetypes.guess_type(image_name)
            if guessed_mime:
                return guessed_mime
        if image_bytes.startswith(b"\xff\xd8\xff"):
            return "image/jpeg"
        if image_bytes.startswith(b"\x89PNG\r\n\x1a\n"):
            return "image/png"
        if image_bytes.startswith(b"RIFF") and b"WEBP" in image_bytes[:16]:
            return "image/webp"
        return None

    @staticmethod
    def _coerce_move_reason(move_plan: MovePlan, execute_move: bool, move_executor: MoveExecutor) -> MovePlan:
        if move_executor == "client" and move_plan.auto_move_allowed:
            return move_plan.model_copy(update={"reason": "client_move_pending"})
        if not execute_move and move_plan.auto_move_allowed:
            return move_plan.model_copy(update={"reason": "execute_move_disabled"})
        return move_plan

    @staticmethod
    def _detect_modality(mime_type: str) -> SourceModality:
        if mime_type in SUPPORTED_IMAGE_TYPES:
            return "image"
        if mime_type in SUPPORTED_AUDIO_TYPES:
            return "audio"
        return "text"

    @staticmethod
    def _resolve_ui_kind(*, document_type: str, source_modality: SourceModality) -> UiDocumentKind:
        if source_modality == "audio":
            return "audio"
        if document_type in {"receipt", "contract", "invoice", "meeting_notes"}:
            return document_type
        return "generic"

    @staticmethod
    def _should_skip_extraction(document_type: str) -> bool:
        return document_type in {"meeting_notes", "generic"}

    @staticmethod
    def _classification_has_empty_core_fields(classification: DocumentClassification) -> bool:
        return not (
            classification.title.strip()
            and classification.summary.strip()
            and classification.template.strip()
        )

    @classmethod
    def _fallback_classification(
        cls,
        *,
        filename: str,
        extracted_text: str,
        source_modality: SourceModality,
        source_path: str | None,
    ) -> DocumentClassification:
        title = cls._derive_fallback_title(filename=filename, source_path=source_path)
        summary_source = cls._sanitize_fallback_summary(extracted_text) or title
        return DocumentClassification(
            document_type="generic",
            template="generic",
            title=title,
            summary=summary_source[:200],
            tags=[],
            language="unknown",
            confidence=0.0,
            ocr_text=summary_source[:200] if source_modality == "image" else None,
            suggested_actions=[],
        )

    @classmethod
    def _derive_fallback_title(cls, *, filename: str, source_path: str | None) -> str:
        candidate = Path(filename).stem.strip() or Path(filename).name.strip()
        if candidate and not cls._is_uuid_like(candidate):
            return candidate

        if source_path:
            source_candidate = Path(source_path).stem.strip()
            source_candidate = cls._strip_staging_prefix(source_candidate)
            if source_candidate and not cls._is_uuid_like(source_candidate):
                return source_candidate

        return "Dokument"

    @classmethod
    def _sanitize_fallback_summary(cls, text: str) -> str:
        sanitized = "".join(char if char.isprintable() or char.isspace() else " " for char in text)
        sanitized = re.sub(r"\s+", " ", sanitized).strip()
        if not sanitized:
            return ""
        words = [
            token
            for token in sanitized.split(" ")
            if token and not cls._is_internal_pipeline_flag(token)
        ]
        return " ".join(words).strip()

    @staticmethod
    def _strip_staging_prefix(name: str) -> str:
        # Tauri staging format: <uuid>-<sanitized_filename>
        return re.sub(
            r"^[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}-",
            "",
            name,
        ).strip()

    @staticmethod
    def _is_uuid_like(value: str) -> bool:
        return bool(
            re.fullmatch(
                r"[0-9a-fA-F]{32}|[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}",
                value,
            )
        )

    @staticmethod
    def _is_internal_pipeline_flag(value: str) -> bool:
        candidate = value.strip().lower()
        return candidate.startswith("classifier_") or candidate.startswith("pdf_") or candidate.endswith("_fallback")

    async def _progress(
        self,
        client_id: str | None,
        request_id: str,
        stage: str,
        message: str,
        data: dict[str, object] | None = None,
    ) -> None:
        payload: dict[str, object] = {
            "type": "job.progress",
            "request_id": request_id,
            "client_id": client_id,
            "stage": stage,
            "message": message,
        }
        if data is not None:
            payload["data"] = data
        await self._emit_event(client_id, payload)

    async def _emit_event(self, client_id: str | None, payload: dict[str, object]) -> None:
        if client_id is None or self.realtime_manager is None:
            return
        self._log_pipeline_event(
            "pipeline.ws.emit",
            request_id=str(payload.get("request_id", "")),
            client_id=client_id,
            event_type=str(payload.get("type", "")),
            stage=str(payload.get("stage", "")) if payload.get("stage") is not None else None,
        )
        await self.realtime_manager.emit_to_client(client_id, payload)

    def _schedule_indexing(
        self,
        *,
        request_id: str,
        client_id: str | None,
        filename: str,
        mime_type: str,
        source_modality: SourceModality,
        record_id: str,
        ui_kind: UiDocumentKind,
        move_status: str,
        indexed_document: IndexedDocument,
    ) -> None:
        task = asyncio.create_task(
            self._index_document_and_finalize(
                request_id=request_id,
                client_id=client_id,
                filename=filename,
                mime_type=mime_type,
                source_modality=source_modality,
                record_id=record_id,
                ui_kind=ui_kind,
                move_status=move_status,
                indexed_document=indexed_document,
            )
        )
        self._background_tasks.add(task)
        task.add_done_callback(self._background_tasks.discard)

    async def _index_document_and_finalize(
        self,
        *,
        request_id: str,
        client_id: str | None,
        filename: str,
        mime_type: str,
        source_modality: SourceModality,
        record_id: str,
        ui_kind: UiDocumentKind,
        move_status: str,
        indexed_document: IndexedDocument,
    ) -> None:
        try:
            await self._progress(client_id, request_id, "indexing", "Indexerar dokument")
            started = time.perf_counter()
            self._log_pipeline_event(
                "pipeline.index.start",
                request_id=request_id,
                client_id=client_id,
                filename=filename,
                mime_type=mime_type,
                source_modality=source_modality,
                record_id=record_id,
            )
            upsert_result = self.search_pipeline.upsert_document(indexed_document)
            if inspect.isawaitable(upsert_result):
                await upsert_result
            elapsed_ms = round((time.perf_counter() - started) * 1000, 2)
            self._log_pipeline_event(
                "pipeline.index.done",
                request_id=request_id,
                client_id=client_id,
                filename=filename,
                mime_type=mime_type,
                source_modality=source_modality,
                record_id=record_id,
                elapsed_ms=elapsed_ms,
            )
        except Exception as error:
            self._log_pipeline_event(
                "pipeline.index.failed",
                request_id=request_id,
                client_id=client_id,
                filename=filename,
                mime_type=mime_type,
                source_modality=source_modality,
                record_id=record_id,
                error=str(error),
            )
            return

        if move_status not in {"auto_pending_client", "awaiting_confirmation"}:
            await self._emit_completed_event(
                client_id=client_id,
                request_id=request_id,
                record_id=record_id,
                ui_kind=ui_kind,
            )

    async def _emit_completed_event(
        self,
        *,
        client_id: str | None,
        request_id: str,
        record_id: str,
        ui_kind: UiDocumentKind,
    ) -> None:
        await self._emit_event(
            client_id,
            {
                "type": "job.completed",
                "request_id": request_id,
                "client_id": client_id,
                "record_id": record_id,
                "ui_kind": ui_kind,
            },
        )

    def _log_pipeline_event(self, event: str, **fields: object) -> None:
        details = " ".join(
            f"{key}={value}"
            for key, value in fields.items()
            if value is not None and value != ""
        )
        logger.info("%s %s", event, details)


__all__ = [
    "ClassificationValidationError",
    "DocumentExtractor",
    "DocumentProcessPipeline",
    "ExtractionValidationError",
    "ExtractionResult",
    "FileOrganizer",
    "IndexedDocument",
    "OllamaServiceError",
    "UnsupportedMediaTypeError",
]
